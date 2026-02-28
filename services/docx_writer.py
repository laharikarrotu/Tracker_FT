from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from services.tailor_resume import TailoredResumeContent


SECTION_MARKERS = {
    "summary": {"summary", "professional summary", "profile"},
    "experience": {"experience", "professional experience", "work experience"},
    "skills": {"skills", "technical skills", "core skills"},
}


@dataclass
class TemplateConstraints:
    summary_count: int
    experience_bullet_count: int


def _normalize(text: str) -> str:
    return " ".join(text.lower().strip().split())


def _is_section_header(text: str, section_key: str) -> bool:
    normalized = _normalize(text)
    return normalized in SECTION_MARKERS[section_key]


def _replace_paragraph_text_preserving_runs(paragraph: Paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _is_bullet_paragraph(paragraph: Paragraph) -> bool:
    style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
    if "bullet" in style_name or "list" in style_name:
        return True
    text = paragraph.text.strip()
    return text.startswith(("•", "-", "*"))


def _find_section_bounds(paragraphs: list[Paragraph], section_key: str) -> tuple[int, int]:
    start_idx = -1
    for idx, paragraph in enumerate(paragraphs):
        if _is_section_header(paragraph.text, section_key):
            start_idx = idx + 1
            break

    if start_idx == -1:
        return -1, -1

    end_idx = len(paragraphs)
    for idx in range(start_idx, len(paragraphs)):
        text = paragraphs[idx].text.strip()
        if not text:
            continue
        normalized = _normalize(text)
        if normalized in SECTION_MARKERS["summary"] | SECTION_MARKERS["experience"] | SECTION_MARKERS["skills"]:
            end_idx = idx
            break
    return start_idx, end_idx


def _find_bullet_paragraphs(paragraphs: list[Paragraph], start_idx: int, end_idx: int) -> list[Paragraph]:
    return [p for p in paragraphs[start_idx:end_idx] if _is_bullet_paragraph(p)]


def read_resume_text(base_resume_path: Path) -> str:
    doc = Document(str(base_resume_path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def analyze_template_constraints(base_resume_path: Path) -> TemplateConstraints:
    doc = Document(str(base_resume_path))
    paragraphs = doc.paragraphs

    summary_start, summary_end = _find_section_bounds(paragraphs, "summary")
    exp_start, exp_end = _find_section_bounds(paragraphs, "experience")

    summary_bullets = (
        _find_bullet_paragraphs(paragraphs, summary_start, summary_end) if summary_start != -1 else []
    )
    exp_bullets = _find_bullet_paragraphs(paragraphs, exp_start, exp_end) if exp_start != -1 else []

    return TemplateConstraints(
        summary_count=max(len(summary_bullets), 3),
        experience_bullet_count=max(len(exp_bullets), 8),
    )


def write_tailored_resume(
    *,
    base_resume_path: Path,
    output_path: Path,
    tailored_content: TailoredResumeContent,
    max_summary_replacements: int | None = None,
    max_experience_replacements: int | None = None,
) -> None:
    doc = Document(str(base_resume_path))
    paragraphs = doc.paragraphs

    summary_start, summary_end = _find_section_bounds(paragraphs, "summary")
    exp_start, exp_end = _find_section_bounds(paragraphs, "experience")
    skills_start, skills_end = _find_section_bounds(paragraphs, "skills")

    if summary_start != -1:
        summary_bullets = _find_bullet_paragraphs(paragraphs, summary_start, summary_end)
        if max_summary_replacements is not None:
            summary_bullets = summary_bullets[:max_summary_replacements]
        for paragraph, line in zip(summary_bullets, tailored_content.summary_points):
            if line:
                _replace_paragraph_text_preserving_runs(paragraph, line)

    if exp_start != -1:
        exp_bullets = _find_bullet_paragraphs(paragraphs, exp_start, exp_end)
        if max_experience_replacements is not None:
            exp_bullets = exp_bullets[:max_experience_replacements]
        for paragraph, line in zip(exp_bullets, tailored_content.experience_points):
            if line:
                _replace_paragraph_text_preserving_runs(paragraph, line)

    if skills_start != -1 and tailored_content.skills_line:
        for paragraph in paragraphs[skills_start:skills_end]:
            if paragraph.text.strip() and not _is_bullet_paragraph(paragraph):
                _replace_paragraph_text_preserving_runs(paragraph, tailored_content.skills_line)
                break

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
